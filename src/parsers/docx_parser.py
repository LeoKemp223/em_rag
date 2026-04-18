"""DOCX 文档解析器"""

from pathlib import Path
from . import DocElement
from .utils import table_to_markdown


class DocxParser:
    def parse(self, file_path: str) -> list[DocElement]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        elements: list[DocElement] = []
        heading_stack: list[str] = []

        for para in doc.paragraphs:
            style = para.style.name or ""
            text = para.text.strip()
            if not text:
                continue

            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                while len(heading_stack) >= level:
                    heading_stack.pop()
                heading_stack.append(text)
                elements.append(DocElement(
                    type="heading", content=text,
                    context_chain=" > ".join(heading_stack), level=level,
                ))
            else:
                elements.append(DocElement(
                    type="text", content=text,
                    context_chain=" > ".join(heading_stack),
                ))

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if len(rows) >= 2:
                md = table_to_markdown(rows)
                if md:
                    elements.append(DocElement(
                        type="table", content=md,
                        context_chain=" > ".join(heading_stack),
                        metadata={"row_count": len(rows)},
                    ))

        return elements
